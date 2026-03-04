import logging

import docx

from app.services.ingestion.parsers.base import FileParser, ParsedContent

logger = logging.getLogger(__name__)


class DocxParser(FileParser):
    EXTENSIONS = {".docx"}

    def parse(self, file_path: str) -> ParsedContent:
        doc = docx.Document(file_path)
        sections: list[str] = []

        body_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if body_paragraphs:
            sections.append("\n\n".join(body_paragraphs))

        header_lines: list[str] = []
        footer_lines: list[str] = []
        for section in doc.sections:
            header_lines.extend(p.text.strip() for p in section.header.paragraphs if p.text.strip())
            footer_lines.extend(p.text.strip() for p in section.footer.paragraphs if p.text.strip())

        if header_lines:
            sections.append("[Header]\n" + "\n".join(header_lines))
        if footer_lines:
            sections.append("[Footer]\n" + "\n".join(footer_lines))

        table_blocks: list[str] = []
        for table_index, table in enumerate(doc.tables, 1):
            row_texts: list[str] = []
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    row_texts.append(" | ".join(cell_texts))
            if row_texts:
                table_blocks.append(f"[Table {table_index}]\n" + "\n".join(row_texts))
        if table_blocks:
            sections.extend(table_blocks)

        skipped_visuals = len(doc.inline_shapes)
        if skipped_visuals > 0:
            logger.warning(
                "DOCX contains %d embedded visual objects not extracted: %s",
                skipped_visuals,
                file_path,
            )

        text = "\n\n".join(section for section in sections if section.strip())
        if not text:
            logger.warning("DOCX file has no text content: %s", file_path)
        return ParsedContent(text=text, language="text")

    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() in self.EXTENSIONS
